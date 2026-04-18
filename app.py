import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import io

# Configuración inicial de la página
st.set_page_config(page_title="Etiquetas Farmacia Word", page_icon="💊")

# --- FUNCIONES DE LIMPIEZA ---
def limpiar_dato(valor):
    if pd.isna(valor) or str(valor).strip().lower() == "nan":
        return ""
    if isinstance(valor, (datetime.datetime, pd.Timestamp)):
        return valor.strftime('%d/%m/%Y')
    val_str = str(valor).strip()
    if " 00:00:00" in val_str:
        val_str = val_str.replace(" 00:00:00", "")
    return val_str

def es_precio(valor):
    try:
        v = str(valor).replace('$', '').strip()
        if not v or "/" in v or "-" in v: return False
        float(v)
        return True
    except: return False

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{border}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '4') 
        node.set(qn('w:color'), '000000')
        tcPr.append(node)

def generar_word(df):
    doc = Document()
    
    # Márgenes de la página
    section = doc.sections[0]
    section.left_margin, section.right_margin = Cm(0.5), Cm(0.5)
    section.top_margin, section.bottom_margin = Cm(1.0), Cm(1.0)

    cols_count = 5
    rows_needed = (len(df) // cols_count) + (1 if len(df) % cols_count != 0 else 0)
    
    table = doc.add_table(rows=rows_needed, cols=cols_count)
    table.autofit = False 
    
    # Configuración de columnas y filas para evitar cortes
    for col in table.columns:
        col.width = Cm(4.0)
    
    for row_obj in table.rows:
        row_obj.height = Cm(2.8)
        # CRÍTICO: Evita que la fila se parta entre dos hojas
        row_obj.allow_break_across_pages = False

    col_med, col_100, col_85 = df.columns[1], df.columns[2], df.columns[3]

    for i, row in df.iterrows():
        r_idx, c_idx = i // cols_count, i % cols_count
        cell = table.cell(r_idx, c_idx)
        set_cell_border(cell)
        
        nombre = str(row[col_med]).strip()
        if nombre.lower() == "nan" or not nombre: continue
        p100, p85 = limpiar_dato(row[col_100]), limpiar_dato(row[col_85])

        # --- NOMBRE DEL PRODUCTO ---
        p_nombre = cell.paragraphs[0]
        p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_n = p_nombre.add_run(nombre)
        run_n.bold = True
        run_n.font.size = Pt(8)
        p_nombre.paragraph_format.space_after = Pt(4)
        
        # --- TABLA INTERNA PARA PRECIOS ---
        sub_table = cell.add_table(rows=2, cols=2)
        sub_table.autofit = True
        
        # Etiquetas: PRECIO / DESCUENTO
        lbl_p = sub_table.cell(0,0).paragraphs[0]
        lbl_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_lp = lbl_p.add_run("PRECIO")
        run_lp.font.size = Pt(5.5)
        
        lbl_d = sub_table.cell(0,1).paragraphs[0]
        lbl_d.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_ld = lbl_d.add_run("DESCUENTO")
        run_ld.font.size = Pt(5.5)
        
        # Valores de Precio
        # Precio Normal
        val_norm = sub_table.cell(1,0).paragraphs[0]
        val_norm.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_v_n = val_norm.add_run(f"${p100}" if es_precio(p100) else p100)
        run_v_n.bold = True
        run_v_n.font.size = Pt(9)
        
        # Precio con Descuento (ROJO)
        val_desc = sub_table.cell(1,1).paragraphs[0]
        val_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        texto_d = f"${p85}" if es_precio(p85) else p85
        run_v_d = val_desc.add_run(texto_d)
        run_v_d.bold = True
        run_v_d.font.color.rgb = RGBColor(200, 0, 0)
        run_v_d.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ ---
st.title("💊 Generador de Etiquetas Profesional (Word)")

with st.expander("📖 ¿Cómo preparar mi archivo?"):
    st.markdown("""
    ### Instrucciones actualizadas:
    1. Las columnas deben ser: **Stock, Nombre, Precio, Descuento**.
    2. Las etiquetas están configuradas para **no cortarse entre hojas**.
    3. Se utiliza el formato **PRECIO / DESCUENTO** con resaltado en rojo.
    """)

archivo = st.file_uploader("Subir archivo de medicamentos", type=["xlsx", "csv"])

if archivo:
    try:
        df = pd.read_csv(archivo) if archivo.name.endswith('.csv') else pd.read_excel(archivo)
        st.success("✨ ¡Archivo cargado!")
        if st.button("🚀 Generar Etiquetas Word"):
            with st.spinner("Generando documento..."):
                word_file = generar_word(df)
                st.download_button(
                    label="📥 Descargar Word",
                    data=word_file,
                    file_name="Etiquetas_Farmacia_Final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except Exception as e:
        st.error(f"Error: {e}")
